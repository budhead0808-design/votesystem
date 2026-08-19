import unittest
import os
import docx
from db import init_db, get_db, generate_voter_code
from gender_rules import calculate_election_results
from word_parser import parse_word_docx

class TestVotingSystem(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        init_db()

    def test_word_parser_enhanced(self):
        """測試強化版 Word 解析器 (表格與段落辨識)"""
        doc = docx.Document()
        doc.add_heading('113學年度教師評審委員會票選名冊', level=1)
        
        # 建立多欄位表格
        table = doc.add_table(rows=1, cols=4)
        table.rows[0].cells[0].text = '號次'
        table.rows[0].cells[1].text = '姓名'
        table.rows[0].cells[2].text = '性別'
        table.rows[0].cells[3].text = '備註'

        r1 = table.add_row().cells
        r1[0].text = '1'
        r1[1].text = '張大明'
        r1[2].text = '男'
        r1[3].text = '國文科 / 兼任行政主任'

        r2 = table.add_row().cells
        r2[0].text = '2'
        r2[1].text = '李美麗'
        r2[2].text = '女'
        r2[3].text = '英文科 / 未兼行政'

        test_docx_path = os.path.join(os.path.dirname(__file__), 'test_sample.docx')
        doc.save(test_docx_path)

        candidates = parse_word_docx(test_docx_path)
        
        # 清理測試檔
        if os.path.exists(test_docx_path):
            os.remove(test_docx_path)

        self.assertEqual(len(candidates), 2)
        self.assertEqual(candidates[0]['name'], '張大明')
        self.assertEqual(candidates[0]['gender'], '男')
        self.assertEqual(candidates[0]['is_admin'], 1)
        self.assertEqual(candidates[1]['name'], '李美麗')
        self.assertEqual(candidates[1]['gender'], '女')
        self.assertEqual(candidates[1]['is_admin'], 0)
        print("[OK] Word Parser Enhanced Test Passed!")

if __name__ == '__main__':
    unittest.main()
