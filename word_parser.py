import docx
import re

# 忽略文字關鍵字 (排除標題與說明)
IGNORE_KEYWORDS = [
    '委員會', '投票', '選票', '候補', '應選', '說明', '圈選', '蓋章', 
    '注意事項', '教育部', '統計', '人數', '日期', '時間', '號次', '姓名', 
    '性別', '備註', '簽章', '兼任行政', '空白'
]

def clean_text(text):
    if not text:
        return ""
    return re.sub(r'[\r\n\t]', '', text).strip()

def parse_word_docx(file_path_or_stream):
    """
    簡化高效版 Word (.docx) 選票與名冊解析器
    聚焦：號次、姓名、性別、兼任行政身分 (不使用/不分學科)
    """
    doc = docx.Document(file_path_or_stream)
    extracted_candidates = []
    seen_names = set()
    current_number = 1

    # 1. 讀取 Word 表格 (Tables)
    if doc.tables:
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                cells = [clean_text(cell.text) for cell in row.cells]
                if not any(cells):
                    continue

                row_str = " ".join(cells)

                # 跳過標頭列
                if row_idx == 0 and any(k in row_str for k in ['號次', '姓名', '選票', '委員']):
                    continue

                name = None
                gender = '未指定'
                is_admin = 0
                num = current_number

                # 識別姓名 (2~4字中文)
                for cell in cells:
                    cell_clean = re.sub(r'[0-9\.\s\(\)（）:：]', '', cell)
                    if 2 <= len(cell_clean) <= 4 and not any(k in cell_clean for k in IGNORE_KEYWORDS):
                        name = cell_clean
                        break

                if not name or name in seen_names:
                    continue

                # 識別號次
                for cell in cells:
                    match_num = re.search(r'(?:號次|編號|號|No\.?)?\s*(\d{1,3})', cell)
                    if match_num and cell != name:
                        try:
                            num = int(match_num.group(1))
                            break
                        except:
                            pass

                # 識別性別
                if '女' in row_str or '女性' in row_str or 'F' in row_str:
                    gender = '女'
                elif '男' in row_str or '男性' in row_str or 'M' in row_str:
                    gender = '男'

                # 識別兼任行政身分
                if any(k in row_str for k in ['兼任行政', '兼行政', '主任', '組長', '秘書', '兼任']):
                    if '未兼' not in row_str and '否' not in row_str:
                        is_admin = 1

                extracted_candidates.append({
                    'candidate_number': num,
                    'name': name,
                    'gender': gender,
                    'is_admin': is_admin
                })
                seen_names.add(name)
                current_number += 1

    # 2. 讀取 Word 段落 (Paragraphs)
    if not extracted_candidates and doc.paragraphs:
        for p in doc.paragraphs:
            text = clean_text(p.text)
            if not text:
                continue

            matches = re.finditer(r'(?:(\d+)[`\.\s、\)\】\:]*)?([\u4e00-\u9fa5]{2,4})(?:[（\(]([^）\)]+)[）\)])?', text)
            for m in matches:
                num_str, name, meta = m.groups()
                if not name or any(k in name for k in IGNORE_KEYWORDS) or name in seen_names:
                    continue

                num = int(num_str) if num_str else current_number
                gender = '未指定'
                is_admin = 0

                search_scope = (text + " " + (meta or ""))
                if '女' in search_scope:
                    gender = '女'
                elif '男' in search_scope:
                    gender = '男'

                if any(k in search_scope for k in ['兼任行政', '兼行政', '主任', '組長', '秘書']):
                    if '未兼' not in search_scope:
                        is_admin = 1

                extracted_candidates.append({
                    'candidate_number': num,
                    'name': name,
                    'gender': gender,
                    'is_admin': is_admin
                })
                seen_names.add(name)
                current_number += 1

    extracted_candidates.sort(key=lambda x: x['candidate_number'])
    return extracted_candidates

if __name__ == '__main__':
    print("Simplified Word parser ready (No department tags).")
