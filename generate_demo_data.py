import os
import docx
from db import init_db, get_db, generate_voter_code

def create_sample_word_docx():
    """建立範例 Word (.docx) 選票檔，供 Word 解析功能測試"""
    doc = docx.Document()
    doc.add_heading('113學年度教師評審委員會 (教評會) 候選人名冊', level=1)
    doc.add_paragraph('請核對以下候選人名冊資料：')

    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '號次'
    hdr_cells[1].text = '姓名'
    hdr_cells[2].text = '性別'
    hdr_cells[3].text = '科別/處室'
    hdr_cells[4].text = '兼任行政身分'

    sample_teachers = [
        ('1', '張哲銘', '男', '國文科', '兼任行政'),
        ('2', '陳雅婷', '女', '英文科', '未兼行政'),
        ('3', '林志豪', '男', '數學科', '未兼行政'),
        ('4', '黃美玲', '女', '國文科', '未兼行政'),
        ('5', '劉柏翰', '男', '自然科', '兼任行政'),
        ('6', '郭靜怡', '女', '社會科', '未兼行政'),
        ('7', '蔡建國', '男', '藝能科', '未兼行政'),
        ('8', '許淑芬', '女', '英文科', '兼任行政'),
        ('9', '鄭家豪', '男', '數學科', '未兼行政'),
        ('10', '謝佩君', '女', '自然科', '未兼行政')
    ]

    for num, name, gender, dept, admin in sample_teachers:
        row_cells = table.add_row().cells
        row_cells[0].text = num
        row_cells[1].text = name
        row_cells[2].text = gender
        row_cells[3].text = dept
        row_cells[4].text = admin

    docx_path = os.path.join(os.path.dirname(__file__), 'demo_ballot.docx')
    doc.save(docx_path)
    print(f"Sample Word file created at: {docx_path}")
    return docx_path

def seed_demo_database():
    """建立模擬委員會與 30 位教師與投票紀錄"""
    init_db()
    conn = get_db()
    cursor = conn.cursor()

    # 1. 建立示範委員會
    cursor.execute("""
        INSERT INTO committees (
            title, description, seats_count, alternate_count, max_votes_per_ballot,
            gender_rule_type, min_male_count, min_female_count, identity_rule_type, min_non_admin_count, status
        ) VALUES (
            '113學年度教師評審委員會 (教評會) 委員票選',
            '依據教育部《高級中等以下學校教師評審委員會設置辦法》規定，任一性別委員應占 1/3 以上。',
            7, 3, 4, 'one_third', 0, 0, 'non_admin_half', 0, 'active'
        )
    """)
    committee_id = cursor.lastrowid

    # 建立另一個性平會投票範例
    cursor.execute("""
        INSERT INTO committees (
            title, description, seats_count, alternate_count, max_votes_per_ballot,
            gender_rule_type, min_male_count, min_female_count, identity_rule_type, min_non_admin_count, status
        ) VALUES (
            '113學年度性別平等教育委員會 (性平會) 票選',
            '依據性平法規定，女性委員不得少於 1/2。',
            5, 2, 3, 'half_female', 0, 0, 'none', 0, 'active'
        )
    """)

    # 2. 建立 20 位教師與候選人名單 (男多女少，便於測試性別保障遞補演算法)
    mock_teachers = [
        ('張哲銘', '男', '國文科', 1),
        ('陳雅婷', '女', '英文科', 0),
        ('林志豪', '男', '數學科', 0),
        ('黃美玲', '女', '國文科', 0),
        ('劉柏翰', '男', '自然科', 1),
        ('郭靜怡', '女', '社會科', 0),
        ('蔡建國', '男', '藝能科', 0),
        ('許淑芬', '女', '英文科', 1),
        ('鄭家豪', '男', '數學科', 0),
        ('謝佩君', '女', '自然科', 0),
        ('賴偉倫', '男', '國文科', 0),
        ('周思婷', '女', '英文科', 0),
        ('楊宗翰', '男', '數學科', 1),
        ('廖佩玲', '女', '社會科', 0),
        ('江冠宇', '男', '自然科', 0),
        ('趙敏君', '女', '藝能科', 0),
        ('潘建廷', '男', '國文科', 1),
        ('邱嘉文', '男', '英文科', 0),
        ('童育德', '男', '數學科', 0),
        ('戴秀英', '女', '自然科', 0)
    ]

    for idx, (name, gender, dept, admin) in enumerate(mock_teachers, start=1):
        vcode = generate_voter_code()
        cursor.execute("INSERT INTO teachers (name, gender, department, is_admin, voter_code) VALUES (?, ?, ?, ?, ?)",
                       (name, gender, dept, admin, vcode))
        
        # 寫入教評會候選人
        cursor.execute("""
            INSERT INTO candidates (committee_id, candidate_number, name, gender, department, is_admin)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (committee_id, idx, name, gender, dept, admin))

    conn.commit()

    # 3. 模擬 15 筆無記名選票投遞 (刻意讓最高票正取中女性偏少，觸發 1/3 性別遞補)
    cursor.execute("SELECT id FROM candidates WHERE committee_id = ?", (committee_id,))
    candidate_ids = [r['id'] for r in cursor.fetchall()]

    # 設定部分高票候選人 (張哲銘、林志豪、劉柏翰、蔡建國、鄭家豪、賴偉倫、楊宗翰 – 男性居多)
    mock_votes = [
        [candidate_ids[0], candidate_ids[2], candidate_ids[4], candidate_ids[6]],
        [candidate_ids[0], candidate_ids[2], candidate_ids[8], candidate_ids[10]],
        [candidate_ids[0], candidate_ids[4], candidate_ids[6], candidate_ids[12]],
        [candidate_ids[2], candidate_ids[8], candidate_ids[10], candidate_ids[1]], # 陳雅婷(女)
        [candidate_ids[0], candidate_ids[2], candidate_ids[4], candidate_ids[6]],
        [candidate_ids[0], candidate_ids[8], candidate_ids[10], candidate_ids[12]],
        [candidate_ids[2], candidate_ids[4], candidate_ids[6], candidate_ids[3]], # 黃美玲(女)
        [candidate_ids[0], candidate_ids[2], candidate_ids[8], candidate_ids[10]],
        [candidate_ids[0], candidate_ids[4], candidate_ids[6], candidate_ids[12]],
        [candidate_ids[2], candidate_ids[8], candidate_ids[10], candidate_ids[5]]  # 郭靜怡(女)
    ]

    for ballot in mock_votes:
        for cid in ballot:
            cursor.execute("INSERT INTO ballots (committee_id, candidate_id) VALUES (?, ?)", (committee_id, cid))

    conn.commit()
    conn.close()
    print("Demo database seeded successfully with mock votes and committees!")

if __name__ == '__main__':
    create_sample_word_docx()
    seed_demo_database()
